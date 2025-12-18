from .llm_client import LLMClient
from .web_scraper import WebScraper
from .google_drive_client import GoogleDriveClient
import os
import re
try:
    from docx import Document
except ImportError:
    Document = None

class ResponseDrafter:
    def __init__(self):
        self.llm = LLMClient()
        self.scraper = WebScraper()
        self.drive_client = GoogleDriveClient()

    def draft_response(self, rfp_text: str, company_url: str = "") -> str:
        """
        Draft a response using company website and source documents from Google Drive.
        """
        website_content = ""
        if company_url:
            website_content = self.scraper.get_website_content(company_url)

        # Get supporting documents from Google Drive Source Information folder
        source_documents = []
        try:
            source_documents = self.drive_client.get_all_rfp_documents()
        except Exception as e:
            print(f"Could not fetch Google Drive source documents: {e}")

        # Build context from source documents
        source_context = ""
        if source_documents:
            source_context = "\n\nSOURCE INFORMATION (from Google Drive):\n"
            for doc in source_documents:
                # Include full content of source documents for better context
                source_context += f"\n--- {doc['name']} ---\n{doc['content']}\n"

        context_section = ""
        if website_content:
            context_section = f"""
            COMPANY KNOWLEDGE BASE (from website {company_url}):
            {website_content}
            """

        prompt = f"""
        You are an expert Proposal Writer. Draft a professional and persuasive response to the following RFP requirement, 
        utilizing the company's knowledge base and source information documents.

        {context_section}
        {source_context}

        RFP REQUIREMENT/TEXT:
        {rfp_text[:5000]}

        INSTRUCTIONS:
        1. Write a response that directly addresses the requirements.
        2. Use information from the source documents to provide specific, accurate details.
        3. Use a professional, confident tone.
        4. Highlight the company's strengths based on the knowledge base and source documents.
        5. If specific details are missing, use descriptive placeholders like [Insert specific metric].
        6. Ensure the response is comprehensive and well-structured.

        RESPONSE:
        """
        
        # Call LLM to generate response
        try:
            response = self.llm.generate_content(prompt)
            return response
        except Exception as e:
            print(f"Error generating LLM response: {e}")
            # Fallback response
            if not source_documents and not company_url:
                return "Please provide a company URL or ensure source documents are available in Google Drive to generate a tailored response."
                
            return f"Based on {len(source_documents)} source documents from Google Drive and website content, here is a draft response:\n\n" \
                   f"Our company is uniquely positioned to deliver this project. We have extensive experience and capabilities " \
                   f"as documented in our source materials. We understand your requirements and propose a comprehensive solution..."

    def generate_draft_document(self, content: str, input_path: str, output_path: str, company_url: str = ""):
        """
        Finds and replaces placeholder text in the document with AI-generated content.
        """
        if not Document or not input_path.endswith('.docx') or not os.path.exists(input_path):
            return None
            
        try:
            doc = Document(input_path)
            
            # Get company context
            website_content = ""
            if company_url:
                website_content = self.scraper.get_website_content(company_url)
            
            # Get source documents from Google Drive
            source_documents = []
            try:
                source_documents = self.drive_client.get_all_rfp_documents()
            except Exception as e:
                print(f"Could not fetch source documents: {e}")
            
            # Find all placeholders in the document
            placeholders = set()
            placeholder_pattern = r'\[([^\]]+)\]'
            
            for paragraph in doc.paragraphs:
                matches = re.findall(placeholder_pattern, paragraph.text)
                placeholders.update(matches)
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            matches = re.findall(placeholder_pattern, paragraph.text)
                            placeholders.update(matches)
            
            # Generate replacements for all placeholders in a single batch using LLM
            # This provides much better context and quality than isolated heuristic checks
            replacements = self._batch_generate_placeholders(placeholders, website_content, source_documents)
            
            # Replace placeholders in paragraphs
            for paragraph in doc.paragraphs:
                for placeholder, replacement in replacements.items():
                    # Handle both with and without brackets keys to be safe
                    target_key = f"[{placeholder}]"
                    if target_key in paragraph.text:
                        paragraph.text = paragraph.text.replace(target_key, replacement)
                    elif placeholder in paragraph.text and not placeholder.startswith("["):
                         # Fallback if the key in dict included brackets already
                         paragraph.text = paragraph.text.replace(placeholder, replacement)

            # Replace placeholders in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for placeholder, replacement in replacements.items():
                                target_key = f"[{placeholder}]"
                                if target_key in paragraph.text:
                                    paragraph.text = paragraph.text.replace(target_key, replacement)

            doc.save(output_path)
            return output_path
            
        except Exception as e:
            print(f"Error modifying docx: {e}")
            return None
    
    def _batch_generate_placeholders(self, placeholders: set, website_content: str, source_documents: list) -> dict:
        """
        Generate content for all placeholders in one go using the LLM and Source Documents.
        Returns a dictionary {placeholder_name: generated_content}
        """
        if not placeholders:
            return {}

        # Prepare Source Context
        source_context = ""
        if source_documents:
            source_context = "SOURCE INFORMATION (from Google Drive):\n"
            for doc in source_documents:
                # Include full content since we are using Gemini 1.5 Flash (1M context)
                content_preview = doc['content']
                
                # DOUBLE SAFETY CHECK for Private Keys in memory
                if "-----BEGIN PRIVATE KEY-----" in content_preview:
                    content_preview = "[SENSITIVE CONTENT REMOVED]"
                
                # Limit size per document safer limit
                if len(content_preview) > 30000:
                    content_preview = content_preview[:30000] + "...[truncated]"
                
                source_context += f"\n--- DOCUMENT: {doc['name']} ---\n{content_preview}\n"
        
        if website_content:
            source_context += f"\n--- WEBSITE CONTENT ---\n{website_content}\n"

        if not source_context:
            return {p: "[No source information available]" for p in placeholders}

        # Prepare Prompt with two-pass strategy
        placeholder_list = "\n".join([f"- {p}" for p in placeholders])
        
        prompt = f"""
        You are an expert data extraction and proposal writing AI. Your task has TWO PHASES:
        
        PHASE 1: EXTRACT ALL INFORMATION FROM SOURCE DOCUMENTS
        Read through ALL the source documents below and create a comprehensive knowledge base. Pay special attention to:
        
        1. TABLES - Extract ALL table data including:
           - Director details tables (Name, Address, Position, Tenure, etc.)
           - Company details tables (Trading name, ABN, ACN, Registration details, etc.)
           - Contact information tables
           - Authority/Authorization tables
        
        2. STRUCTURED DATA - Look for:
           - Names (directors, officers, authorized persons)
           - Numbers (ABN, ACN, phone numbers)
           - Addresses (registered office, principal place of business)
           - Dates (incorporation date, tenure periods)
           - Yes/No answers (bankruptcy declarations, etc.)
        
        3. CONTEXT - Understand the relationship between data:
           - Which address belongs to which director
           - Which position each person holds
           - How long each director has served
        
        PHASE 2: FILL PLACEHOLDERS INTELLIGENTLY
        Now use your extracted knowledge to fill these placeholders:
        
        {placeholder_list}
        
        MATCHING RULES:
        - "Name" in a Directors table → Extract director names from your knowledge base
        - "Address" in a Directors table → Extract corresponding director addresses
        - "Position held" → Extract director positions/roles
        - "Length of tenure" → Extract tenure information
        - "Trading name" → Extract company trading name
        - "ABN" or "Australian Business Number" → Extract ABN number
        - "ACN" or "Australian Company Number" → Extract ACN number
        - "Registered office" or "Address of registered office" → Extract registered address
        - "Date of incorporation" → Extract incorporation date
        - "Name of authorised officer" → Extract authorized officer name
        - Bankruptcy/administration questions → Extract Yes/No answers
        
        CRITICAL INSTRUCTIONS:
        1. If you found information in PHASE 1, USE IT EXACTLY - don't say "not available"
        2. For table cells, extract the SPECIFIC value for that row/column combination
        3. If a placeholder asks for "Name" and you're in row 1 of Directors table, use Director 1's name
        4. If a placeholder asks for "Address" and you're in row 2 of Directors table, use Director 2's address
        5. Match the context - if the placeholder is in "Table 5. Directors' details", use director information
        6. If truly not found after thorough search, use "[Information not available in source documents]"
        
        SOURCE DOCUMENTS:
        {source_context}
        
        OUTPUT FORMAT:
        Return a JSON object where:
        - Keys = exact placeholder names from the list above
        - Values = the extracted information or generated content
        
        EXAMPLE OUTPUT STRUCTURE:
        {{
            "Trading name": "Intelia Pty Ltd",
            "Name": "John Smith",
            "Address": "123 Main St, Sydney NSW 2000",
            "Position held": "Director",
            "Length of tenure": "5 years"
        }}
        
        Return ONLY the JSON object. No markdown, no explanations, no extra text.
        """

        try:
            print(f"Generating batch response for {len(placeholders)} placeholders...")
            response_text = self.llm.generate_content(prompt)
            
            # Check if LLM client returned an error string
            if response_text.startswith("Error generating content"):
                print(f"LLM Generation Error: {response_text}")
                return {p: f"[{response_text}]" for p in placeholders}

            # Robust JSON extraction
            cleaned_replacements = {}
            import json
            import re
            
            # Find the first '{' and last '}' to extract the JSON object
            json_match = re.search(r'(\{.*\})', response_text, re.DOTALL)
            
            if json_match:
                json_str = json_match.group(1)
                try:
                    replacements = json.loads(json_str)
                    
                    # Validate keys match
                    for p in placeholders:
                        # the LLM might return keys with brackets or w/o, try to match
                        if p in replacements:
                            cleaned_replacements[p] = replacements[p]
                        elif f"[{p}]" in replacements:
                            cleaned_replacements[p] = replacements[f"[{p}]"]
                        else:
                            # Fallback if missed
                            cleaned_replacements[p] = "[Information not found in knowledge base]"
                            
                    return cleaned_replacements
                    
                except json.JSONDecodeError as je:
                    print(f"JSON Parse Error: {je}. Response was: {response_text[:200]}...")
                    return {p: "[Error: AI response was not valid data]" for p in placeholders}
            else:
                print(f"No JSON found in response. Response was: {response_text[:200]}...")
                return {p: "[Error: AI provided no structured data]" for p in placeholders}

        except Exception as e:
            print(f"Error in batch generation: {e}")
            return {p: f"[System Error: {str(e)}]" for p in placeholders}



